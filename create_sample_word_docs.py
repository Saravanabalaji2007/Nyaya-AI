"""
NYAYA AI 3.0 – Sample Legal Documents Generator
Generates Word Documents (.docx), Plain Text (.txt), and PDF (.pdf) for:
- Rental Agreement (Safe, Medium Risk, High Risk)
- Insurance Policy (Safe, Medium Risk, High Risk)
- Loan Document (Safe, Medium Risk, High Risk)
"""

import os
import sys

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8')

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sample_docs")
os.makedirs(OUT_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Document Data Definitions
# ---------------------------------------------------------------------------
DOCUMENTS = {
    # ==================== 1. RENTAL AGREEMENTS ====================
    "rental_agreement_safe": {
        "title": "STANDARD RESIDENTIAL LEASE AGREEMENT (FAIR & SAFE)",
        "risk_level": "🟢 Safe (Low Risk ~10%)",
        "description": "Compliant with Tamil Nadu Regulation of Rights and Responsibilities of Landlords and Tenants Act 2017.",
        "paragraphs": [
            "THIS RESIDENTIAL LEASE AGREEMENT is executed on this 1st day of September, 2026 at Chennai, Tamil Nadu.",
            "BETWEEN: Mr. S. Raghavan (hereinafter referred to as the 'Lessor/Owner') AND Mr. K. Vijay (hereinafter referred to as the 'Lessee/Tenant').",
            "",
            "1. LEASE TERM & RENT:",
            "The lease shall be valid for a tenure of 11 months commencing from September 1, 2026. The agreed monthly rent is Rs. 20,000 (Twenty Thousand Rupees only), payable on or before the 5th of each calendar month. Any annual revision of rent shall not exceed 5% upon mutual written agreement.",
            "",
            "2. REFUNDABLE SECURITY DEPOSIT:",
            "The Lessee has deposited a refundable security deposit of Rs. 60,000 (equivalent to 3 months' rent). The entire security deposit shall be refunded in full within 15 days of vacation of premises, subject only to actual unpaid utility bills and physical structural damage.",
            "",
            "3. MAINTENANCE & STRUCTURAL REPAIRS:",
            "The Lessor shall remain strictly responsible for major structural maintenance, waterproofing, and external repairs. The Lessee shall bear regular electricity and domestic water consumption charges.",
            "",
            "4. TERMINATION & NOTICE PERIOD:",
            "Either party may terminate this lease by serving a 30-day prior written notice. No unilateral termination penalties or arbitrary forfeitures shall apply.",
            "",
            "5. DISPUTE RESOLUTION:",
            "All disputes shall be subject to the jurisdiction of the Rent Court / Rent Tribunal established under the Tamil Nadu Tenancy Act 2017 in Chennai."
        ]
    },
    "rental_agreement_medium_risk": {
        "title": "RESIDENTIAL TENANCY AGREEMENT (MODERATE RISK)",
        "risk_level": "🟡 Medium Risk (~45%)",
        "description": "Contains mandatory painting deductions, elevated notice periods, and rent escalation clauses.",
        "paragraphs": [
            "THIS RESIDENTIAL TENANCY AGREEMENT is executed on August 15, 2026 at Coimbatore, Tamil Nadu.",
            "BETWEEN: Southern Properties & Leasing LLP (Lessor) AND Mr. Arvind Swaminathan (Lessee).",
            "",
            "1. TERM & AUTOMATIC RENT ESCALATION:",
            "The lease tenure is 11 months. The monthly rent is Rs. 25,000. Rent shall automatically increase by 10% upon completion of 11 months without further notice or negotiation.",
            "",
            "2. SECURITY DEPOSIT & MANDATORY DEDUCTIONS:",
            "The Lessee shall deposit Rs. 1,50,000 as security deposit. Upon vacating, a mandatory fixed non-negotiable deduction of Rs. 25,000 will be made towards repainting and deep-cleaning, irrespective of the physical condition of the premises.",
            "",
            "3. LATE PAYMENT PENALTY:",
            "Rent received after the 5th of the month will attract a late payment fee of Rs. 500 per day until realization.",
            "",
            "4. TERMINATION NOTICE:",
            "The Lessee must provide a mandatory 3-month written notice period before vacating. Early termination by the tenant will result in forfeiture of 50% of the security deposit as termination fees.",
            "",
            "5. JURISDICTION:",
            "Any disputes shall be resolved in the local Civil Courts of Coimbatore, Tamil Nadu."
        ]
    },
    "rental_agreement_high_risk": {
        "title": "COMMERCIAL & RESIDENTIAL LEASE CONTRACT (HIGH RISK / TRAP CLAUSES)",
        "risk_level": "🔴 High Risk (~85%)",
        "description": "Contains unconscionable lock-in penalties, hidden charges, compounding daily interest, and waiver of legal rights.",
        "paragraphs": [
            "THIS BINDING LEASE & SERVICES CONTRACT is entered into on August 20, 2026 at Chennai.",
            "BY AND BETWEEN: Apex Realty Assets Management Pvt. Ltd. (the 'Company') AND the Signatory/Tenant.",
            "",
            "1. FEES, HIDDEN CHARGES & COMPOUNDING PENALTIES:",
            "The monthly occupancy fee is Rs. 35,000. The Company reserves the absolute unilateral right to levy hidden charges, maintenance fees, and administrative charges without prior notification. Delayed payments beyond 24 hours incur an immediate compounding penalty of 10% per day on the total outstanding balance.",
            "",
            "2. LOCK-IN PERIOD & TOTAL DEPOSIT FORFEITURE:",
            "A mandatory lock-in period of 24 months applies. In the event of early termination, the entire security deposit of Rs. 2,00,000 shall be fully non-refundable, and the Tenant shall remain liable to pay the full rent for the remainder of the 24-month term.",
            "",
            "3. AUTOMATIC RENEWAL & UNILATERAL RENT HIKE:",
            "This contract will automatically renew for subsequent 3-year terms with an automatic 20% annual rent increase unless cancelled by registered post 180 days prior.",
            "",
            "4. INDEMNITY & WAIVER OF STATUTORY RIGHTS:",
            "The Tenant agrees to indemnify and hold harmless the Company against all liabilities and damages. The Tenant explicitly agrees to waive statutory rights to approach the Rent Court or Consumer Forum.",
            "",
            "5. BINDING ARBITRATION:",
            "All disputes shall be settled exclusively by sole binding arbitration appointed unilaterally by the Company at New Delhi, and civil court jurisdiction is explicitly barred."
        ]
    },

    # ==================== 2. INSURANCE POLICIES ====================
    "insurance_policy_safe": {
        "title": "COMPREHENSIVE HEALTH & LIFE INSURANCE POLICY (IRDAI COMPLIANT)",
        "risk_level": "🟢 Safe (Low Risk ~12%)",
        "description": "Standard IRDAI consumer-friendly policy with clear claim terms, 15-day free-look period, and Ombudsman protection.",
        "paragraphs": [
            "POLICY SCHEDULE & TERMS – NYAYA SURAKSHA HEALTH SHIELD (POLICY NO: IRDAI/2026/HL/8821)",
            "INSURER: National Health & General Insurance Co. Ltd. | INSURED: Mr. Anand Sundaram | SUM INSURED: Rs. 10,00,000",
            "",
            "1. COVERAGE & BENEFITS:",
            "This policy provides comprehensive coverage for hospitalization, day-care procedures, pre-hospitalization (60 days), and post-hospitalization (90 days) medical expenses without room-rent capping.",
            "",
            "2. FREE LOOK PERIOD & CANCELLATION REFUND:",
            "The Insured is entitled to a 15-day Free Look Period from the date of receipt of this policy. If cancelled within this period, the full premium shall be refunded after deducting only medical examination and stamp duty charges.",
            "",
            "3. GRACE PERIOD FOR PREMIUM:",
            "A statutory grace period of 30 days is provided for yearly renewal premium payment. Coverage continuity benefits shall not be forfeited during the grace period.",
            "",
            "4. PRE-EXISTING DISEASE WAITING PERIOD:",
            "Pre-existing ailments declared at inception shall be covered after a standard waiting period of 24 months of continuous policy renewals.",
            "",
            "5. GRIEVANCE REDRESSAL & INSURANCE OMBUDSMAN:",
            "In case of claim repudiation or dispute, the Insured has the statutory right to escalate complaints to the Insurance Ombudsman under Insurance Ombudsman Rules 2017 or file before the District Consumer Commission."
        ]
    },
    "insurance_policy_medium_risk": {
        "title": "STANDARD HEALTH INSURANCE POLICY (MODERATE CO-PAY & SUB-LIMITS)",
        "risk_level": "🟡 Medium Risk (~48%)",
        "description": "Contains mandatory 20% co-payment, ICU sub-limits, and extended waiting periods.",
        "paragraphs": [
            "POLICY SCHEDULE – PRUDENT HEALTH PROTECT (POLICY NO: PHP/2026/5541)",
            "INSURER: Starline General Insurance Ltd. | INSURED: Ms. Deepa Nambiar | SUM INSURED: Rs. 5,00,000",
            "",
            "1. CO-PAYMENT CLAUSE:",
            "A mandatory 20% co-payment shall apply to all claims incurred in non-network hospitals and for insured persons above 50 years of age.",
            "",
            "2. ROOM RENT & ICU SUB-LIMITS:",
            "Room rent is strictly capped at 1% of the Sum Insured (Rs. 5,000/day) and ICU charges at 2% (Rs. 10,000/day). If the Insured opts for a room with higher tariff, proportionate deductions will be applied across all medical bills.",
            "",
            "3. EXTENDED PRE-EXISTING AILMENT WAITING PERIOD:",
            "Pre-existing illnesses (hypertension, diabetes, joint disorders) shall be subject to a strict 48-month continuous waiting period before coverage commences.",
            "",
            "4. CANCELLATION REFUND GRID:",
            "Upon policy cancellation after the 30th day, refund of premium shall be retained on short-period scale: 50% retained up to 3 months, and 100% premium retained after 6 months.",
            "",
            "5. DISPUTE RESOLUTION:",
            "Disputes shall be handled through the internal Grievance Officer and subsequent Civil Courts."
        ]
    },
    "insurance_policy_high_risk": {
        "title": "PRIVATE EXTENDED WARRANTY & LIFE-HEALTH COVER (HIGH RISK / ABUSIVE EXCLUSIONS)",
        "risk_level": "🔴 High Risk (~90%)",
        "description": "Contains unilateral claim rejection rights, total non-refundable premium, exclusion of critical conditions, and waiver of Insurance Ombudsman.",
        "paragraphs": [
            "CERTIFICATE OF INSURANCE & CONTRACT – GLOBAL ASSURE ULTRA (POLICY NO: GAU/2026/9901)",
            "UNDERWRITER: Horizon Global Reassurance Consortium Ltd. | INSURED: Applicant Signatory",
            "",
            "1. UNILATERAL CLAIM REPUDIATION & PROOF OF LOSS:",
            "The Company reserves the sole, unchallengeable discretion to deny and reject any claim without assigning detailed statutory justification if internal medical assessors deem the treatment non-essential.",
            "",
            "2. 100% NON-REFUNDABLE PREMIUMS & NO FREE LOOK:",
            "All paid premiums, processing charges, and policy fees are 100% non-refundable under all circumstances from the exact hour of policy issuance. Free-look cancellation is expressly excluded.",
            "",
            "3. CRITICAL ILLNESS PERMANENT EXCLUSIONS:",
            "Heart conditions, oncological treatments, neurological conditions, respiratory disorders, and sudden lifestyle illnesses are permanently excluded from claim eligibility.",
            "",
            "4. WAIVER OF STATUTORY OMBUDSMAN & CONSUMER REMEDIES:",
            "The Insured explicitly agrees to forfeit the right to approach the Insurance Ombudsman or Consumer Disputes Redressal Commission.",
            "",
            "5. SOLE PRIVATE ARBITRATION IN MAURITIUS / OVERSEAS:",
            "Any claim dispute shall be referred exclusively to private commercial arbitration conducted in Singapore or Mauritius at the shared expense of the claimant, barring Indian court jurisdiction."
        ]
    },

    # ==================== 3. LOAN DOCUMENTS ====================
    "loan_agreement_safe": {
        "title": "HOME LOAN AGREEMENT (FAIR & RBI COMPLIANT)",
        "risk_level": "🟢 Safe (Low Risk ~15%)",
        "description": "Fully compliant with Reserve Bank of India (RBI) Fair Practices Code, zero foreclosure penalty on floating rates.",
        "paragraphs": [
            "HOUSING FINANCE & LOAN AGREEMENT (LOAN ACC: HL/2026/4401)",
            "BETWEEN: State Housing Finance Bank Ltd. (the 'Lender') AND Mr. Rajesh Kanna & Mrs. Shanthi Kanna (the 'Borrower').",
            "",
            "1. SANCTIONED LOAN AMOUNT & INTEREST RATE:",
            "The Lender sanctions a Home Loan of Rs. 45,00,000 (Rupees Forty-Five Lakhs only) at an external benchmark linked floating interest rate of 8.40% per annum (EBLR + 0.90%). The repayment schedule is fixed for 240 equal monthly installments (EMI) of Rs. 38,765.",
            "",
            "2. ZERO FORECLOSURE / PREPAYMENT CHARGES:",
            "In strict compliance with RBI circulars, NO foreclosure fees, prepayment penalties, or hidden exit charges shall be levied on partial or full prepayment of the floating-rate home loan.",
            "",
            "3. TRANSPARENT FEE STRUCTURE:",
            "A one-time administrative processing fee of Rs. 10,000 + GST has been charged upfront. No additional unnotified annual inspection fees or ledger fees will be deducted.",
            "",
            "4. DEFAULT & FAIR NOTICE PROCEDURE:",
            "In case of payment delay, penal interest is capped at 2% per annum on the overdue amount only. The Lender shall issue a 60-day statutory notice before initiating any SARFAESI recovery proceedings.",
            "",
            "5. JURISDICTION & GRIEVANCE REDRESSAL:",
            "Borrowers may escalate unresolved grievances to the RBI Banking Ombudsman. Disputes remain subject to the Civil Courts in Chennai."
        ]
    },
    "loan_agreement_medium_risk": {
        "title": "PERSONAL & VEHICLE LOAN AGREEMENT (MODERATE CHARGES & SPREAD RESET)",
        "risk_level": "🟡 Medium Risk (~52%)",
        "description": "Contains 2.5% processing fee, 24% annual penal interest, and lender spread reset rights.",
        "paragraphs": [
            "RETAIL TERM LOAN CONTRACT (CONTRACT NO: RL/2026/7723)",
            "BETWEEN: QuickCredit Commercial Finance Ltd. (NBFC) AND Mr. Karthikeyan Murugan (Borrower).",
            "",
            "1. LOAN AMOUNT & MARGIN SPREAD:",
            "Sanctioned principal: Rs. 6,00,000 at 12.5% interest per annum for a 36-month tenure. The Lender reserves the right to reset the benchmark spread annually based on market liquidity.",
            "",
            "2. PROCESSING & ADMINISTRATIVE LEVIES:",
            "Upfront processing deduction of 2.5% (Rs. 15,000) + documentation fee of Rs. 3,500 shall be deducted directly from the disbursal sum.",
            "",
            "3. LATE EMI PENAL CHARGES:",
            "Overdue installments will incur a penal interest of 24% per annum (2% per month) plus bounce charges of Rs. 750 per dishonored NACH mandate.",
            "",
            "4. PRE-CLOSURE CHARGES:",
            "Foreclosure of the loan within 12 months is subject to a 3% pre-closure fee on the outstanding principal balance.",
            "",
            "5. JURISDICTION:",
            "Governed by the laws of India with jurisdiction in the competent courts of Coimbatore, Tamil Nadu."
        ]
    },
    "loan_agreement_high_risk": {
        "title": "INSTANT DIGITAL MICRO-LOAN CONTRACT (HIGH RISK / PREDATORY TERMS)",
        "risk_level": "🔴 High Risk (~95%)",
        "description": "Contains 36% compounding interest, irrevocable salary debit, 5% exit fees, unlimited indemnity, and private arbitration.",
        "paragraphs": [
            "DIGITAL CREDIT LINE & UNSECURED LOAN CONTRACT (REF NO: DCL/2026/1099)",
            "LENDER: TurboCash Capital Online FinTech Pvt. Ltd. | BORROWER: Digital Applicant / Signatory",
            "",
            "1. COMPOUNDING PENAL INTEREST & HIDDEN PLATFORM CHARGES:",
            "Loan Principal: Rs. 1,00,000. Interest Rate: 28% APR. The Lender reserves the right to deduct unannounced monthly technology fees and loan servicing charges. Any single missed EMI results in an immediate compounding penalty of 36% per annum calculated on a daily compounding basis.",
            "",
            "2. IRREVOCABLE POWER OF ATTORNEY & SALARY LIEN:",
            "The Borrower hereby executes an irrevocable Power of Attorney granting the Lender direct authority to deduct payments directly from the Borrower's bank account, employer payroll, and secondary deposit accounts without prior intimation.",
            "",
            "3. PREDATORY FORECLOSURE & TERMINATION PENALTY:",
            "Prepayment or early closure incurs a mandatory 5% exit penalty on the gross original loan sanctioned amount, plus forfeiture of all unadjusted interest.",
            "",
            "4. UNILATERAL TERMS MODIFICATION & INDEMNITY:",
            "The Lender reserves the absolute right to alter the interest rate, tenure, and fee schedule at its sole discretion. The Borrower agrees to indemnify and hold harmless the Lender against all claims and damages.",
            "",
            "5. WAIVER OF COURT JURISDICTION & SOLE ARBITRATION:",
            "The Borrower expressly agrees to waive all statutory rights to file proceedings in the Consumer Forum or Civil Court. All disputes shall be resolved exclusively via binding private arbitration conducted solely by an arbitrator selected by the Lender in Mumbai."
        ]
    }
}

# ---------------------------------------------------------------------------
# Generators
# ---------------------------------------------------------------------------

def generate_word_doc(key: str, data: dict):
    """Generates formatted Microsoft Word (.docx) document."""
    filepath = os.path.join(OUT_DIR, f"{key}.docx")
    doc = Document()
    
    # Title
    p_title = doc.add_paragraph()
    run_title = p_title.add_run(data["title"])
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.name = "Arial"
    if "Safe" in data["risk_level"]:
        run_title.font.color.rgb = RGBColor(16, 185, 129)
    elif "Medium" in data["risk_level"]:
        run_title.font.color.rgb = RGBColor(245, 158, 11)
    else:
        run_title.font.color.rgb = RGBColor(239, 68, 68)
        
    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run(f"Risk Rating: {data['risk_level']} | NYAYA AI 3.0 Test Document\n{data['description']}")
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    
    doc.add_paragraph("―" * 50)
    
    for text in data["paragraphs"]:
        if text.strip().startswith(("1.", "2.", "3.", "4.", "5.")):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = "Arial"
        elif text == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.size = Pt(10.5)
            r.font.name = "Arial"
            
    doc.save(filepath)
    print(f"  ✓ Word Document generated: {key}.docx")


def generate_txt_doc(key: str, data: dict):
    """Generates plain text (.txt) document."""
    filepath = os.path.join(OUT_DIR, f"{key}.txt")
    lines = [
        data["title"],
        f"Risk Rating: {data['risk_level']} | NYAYA AI 3.0 Benchmark Document",
        data["description"],
        "=" * 70,
        ""
    ] + data["paragraphs"]
    
    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"  ✓ Text Document generated: {key}.txt")


def generate_pdf_doc(key: str, data: dict):
    """Generates formatted PDF document using ReportLab."""
    filepath = os.path.join(OUT_DIR, f"{key}.pdf")
    doc = SimpleDocTemplate(filepath, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    
    from reportlab.lib import colors
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#0f172a')
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor('#334155')
    )
    heading_style = ParagraphStyle(
        'DocHeading',
        parent=styles['Heading2'],
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor('#1e3a8a')
    )
    
    story = [
        Paragraph(f"<b>{data['title']}</b>", title_style),
        Spacer(1, 4),
        Paragraph(f"<i>Risk Rating: {data['risk_level']} – NYAYA AI 3.0 Test Document</i>", body_style),
        Spacer(1, 10),
    ]
    
    for text in data["paragraphs"]:
        if text.strip().startswith(("1.", "2.", "3.", "4.", "5.")):
            story.append(Spacer(1, 6))
            story.append(Paragraph(f"<b>{text}</b>", heading_style))
            story.append(Spacer(1, 3))
        elif text == "":
            story.append(Spacer(1, 4))
        else:
            story.append(Paragraph(text, body_style))
            story.append(Spacer(1, 4))
            
    doc.build(story)
    print(f"  ✓ PDF Document generated: {key}.pdf")


def main():
    print("=" * 65)
    print("  Generating 9 Comprehensive Word (.docx), TXT, and PDF Documents")
    print("=" * 65)
    for key, data in DOCUMENTS.items():
        print(f"\nProcessing [{key}]...")
        generate_word_doc(key, data)
        generate_txt_doc(key, data)
        generate_pdf_doc(key, data)
        
    print("\n" + "=" * 65)
    print(f"  🎉 All 27 sample files (.docx, .txt, .pdf) successfully generated in:\n  {OUT_DIR}")
    print("=" * 65)

if __name__ == "__main__":
    main()
